#PDF
#从PDF提取文本
    # import pypdf
    # pdf=r"D:\python\project\FunctionTest\处理PDF和Word文档测试目录\combinedminutes.pdf"
    # pdfFileobj=open(pdf,'rb')
    #
    # #获取pdf的页数
    # pdfReader=pypdf.PdfReader(pdfFileobj)
    # print(pdfReader.get_num_pages())
    #
    # pageobj=pdfReader.get_page(0)
    # print(pageobj.extract_text())
    # pdfFileobj.close()

#解密PDF
import pypdf
pdf=r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\encrypted.pdf'
pdfFileobj=open(pdf,'rb')

pdfReader=pypdf.PdfReader(pdfFileobj)
print(pdfReader.is_encrypted)
#不解密测试
try:
    print(pdfReader.get_page(0))
except pypdf.errors.FileNotDecryptedError:
    print('文件需要解密')

print(pdfReader.decrypt('rosebud'))
print(pdfReader.get_page(0))

# 创建PDF
    #复制页面
        # import pypdf
        # pdf1=r"D:\python\project\FunctionTest\处理PDF和Word文档测试目录\meetingminutes.pdf"
        # pdf2=r"D:\python\project\FunctionTest\处理PDF和Word文档测试目录\meetingminutes2.pdf"
        # pdf1Fileobj=open(pdf1,'rb')
        # pdf2Fileobj=open(pdf2,'rb')
        #
        # pdf1Reader=pypdf.PdfReader(pdf1Fileobj)
        # pdf2Reader=pypdf.PdfReader(pdf2Fileobj)
        # pdfWriter=pypdf.PdfWriter()
        #
        # for pageNum in range(pdf1Reader.get_num_pages()):
        #     pageobj=pdf2Reader.get_page(pageNum)
        #     pdfWriter.add_page(pageobj)
        #
        # for pageNum in range(pdf2Reader.get_num_pages()):
        #     pageobj=pdf2Reader.get_page(pageNum)
        #     pdfWriter.add_page(pageobj)
        #
        # pdfOutputFile=open(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\meetingminutes.pdf','wb')
        # pdfWriter.write(pdfOutputFile)
        #
        # pdfOutputFile.close()
        # pdf1Fileobj.close()
        # pdf2Fileobj.close()

    #旋转页面
        # import pypdf
        # minuterFiles=open(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\meetingminutes.pdf','rb')
        # pdfReader=pypdf.PdfReader(minuterFiles)
        # page=pdfReader.get_page(0)
        # page.rotate(90)
        #
        # pdfWrite=pypdf.PdfWriter()
        # pdfWrite.add_page(page)
        #
        # resultpdfFile=open(r"D:\python\project\FunctionTest\处理PDF和Word文档测试目录\meetingminutes.pdf",'wb')
        # pdfWrite.write(resultpdfFile)
        #
        # resultpdfFile.close()
        # minuterFiles.close()
    #叠加页面
        # import pypdf
        # pdf1File=open(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\meetingminutes.pdf','rb')
        # pdfReader=pypdf.PdfReader(pdf1File)
        # page1=pdfReader.get_page(0)
        #
        # pdf2File=open(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\watermark.pdf','rb')
        # pdfwateReader=pypdf.PdfReader(pdf2File)
        # page2=pdfwateReader.get_page(0)
        #
        # page1.merge_page(page2)
        #
        # pdfWrite=pypdf.PdfWriter()
        # pdfWrite.add_page(page1)
        #
        # for pdfpage in range(1,len(pdfReader.pages)):
        #     pdfobj=pdfReader.get_page(pdfpage)
        #     pdfWrite.add_page(pdfobj)
        #
        # pdfOutput=open(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\水印.pdf','wb')
        # pdfWrite.write(pdfOutput)
        # pdf1File.close()
        # pdf2File.close()

    #加密PDF
        # import pypdf
        # pdfFile=open(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\meetingminutes.pdf','rb')
        #
        # pdfReader=pypdf.PdfReader(pdfFile)
        # pdfWriter=pypdf.PdfWriter()
        #
        # for pageNum in range(pdfReader.get_num_pages()):
        #     pdfWriter.add_page(pdfReader.get_page(pageNum))
        ##对PDF进行加密
        # pdfWriter.encrypt('123')
        #
        # resultPdf=open(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\加密.pdf','wb')
        # pdfWriter.write(resultPdf)
        # resultPdf.close()

#项目：从多个PDF中合并选择的页面 【略】

#Word文档
    #读取Word文档
        # import docx
        # doc=docx.Document(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\demo.docx')
        # print(type(doc))
        # print(len(doc.paragraphs))
        # print(doc.paragraphs[0].text)
        # print(doc.paragraphs[1].text)
        # print(len(doc.paragraphs[1].runs))
        # print(doc.paragraphs[1].runs[0].text)
        # print(doc.paragraphs[1].runs[1].text)
        # print(doc.paragraphs[1].runs[4].text)

    #从.docx文档中取得完整的文档
        # import docx
        # def getText(filename):
        #     doc=docx.Document(filename)
        #     fullText=[]
        #     for para in doc.paragraphs:
        #         fullText.append(' '+para.text)
        #     return '\n'.join(fullText)
        #
        # print(getText(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\demo.docx'))

    #设置 Paragraph 和 Run 对象的样式 【Run添加样式:Runobj.style='quoteChar'需要加'Char' Paragraph添加样式：paragraphobj.style='quote'】

    #创建带有非默认样式的 Word 文档 【使用新曾样式时需要打开word手动添加样式之后才能使用python进行样式编辑】

    #Run 属性
        # import docx
        # doc=docx.Document(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\demo.docx')
        #
        # #查看字体样式
        # print(doc.paragraphs[0].text)
        # print(doc.paragraphs[0].style)
        # #改变段落样式
        # doc.paragraphs[0].style='Normal'
        # print(doc.paragraphs[0].style)
        #
        # print(doc.paragraphs[1].text)
        # print(doc.paragraphs[1].runs[0].text)
        # print(doc.paragraphs[1].runs[1].text)
        # print(doc.paragraphs[1].runs[2].text)
        # print(doc.paragraphs[1].runs[3].text)
        # #改变runs样式
        # doc.paragraphs[1].runs[0].style='QuoteChar'
        # doc.paragraphs[1].runs[1].underline=True
        # doc.paragraphs[1].runs[3].underline=True
        #
        # doc.save(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\demo.docx')

    #写入 Word 文档 [add_paragraph add_run]
        # import docx
        # doc=docx.Document()
        # doc.add_paragraph('你好世界！')
        # #新增内容（以段为单位）
        # paraObj1=doc.add_paragraph('这是一个残酷的世界')
        # paraObj2=doc.add_paragraph('这个世界并不友好')
        # #追加内容
        # paraObj1.add_run('这里将内容添加进去')
        #
        # doc.add_paragraph("向世界说你好！",'Title')
        # doc.save(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\New_word.docx')

    #添加标题 [add_deading]
        # import docx
        # doc=docx.Document()
        # doc.add_heading('头部信息-0',0)
        # doc.add_heading('头部信息-1',1)
        # doc.add_heading('头部信息-2',2)
        # doc.save(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\标题.docx')

    #添加换行符和换页符 [add_break]
        # import docx
        # doc=docx.Document()
        # doc.add_paragraph('这是第一个段落~')
        # doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        # doc.add_paragraph('这是第二个段落~')
        # doc.save(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\换行符.docx')

    #添加图像 [add_picture]
        # import docx
        # doc=docx.Document()
        # doc.add_heading('这个一个一级抬头',0)
        # doc.add_paragraph('这是第一个段落')
        # # doc.add_picture(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\pph.png',width=docx.shared.Inches(1),height=docx.shared.Cm(4))
        # doc.add_picture(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\pph.png')
        # doc.save(r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\添加图片.docx')

#从Word文档中创建PDF
import win32com.client
import docx
wordFilename=r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\word_for_pdf.docx'
pdfFilename=r'D:\python\project\FunctionTest\处理PDF和Word文档测试目录\pdf.pdf'

#创建一个docx文件
doc=docx.Document()
doc.add_paragraph('这是一个段落')
doc.save(wordFilename)

#转换为PDF
wdFormatPDF=17
wordObj=win32com.client.Dispatch('Word.Application')

docObj=wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename,FileFormat=wdFormatPDF)
docObj.Close()

#